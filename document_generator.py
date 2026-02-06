"""
PDF and Document Generation Utilities
Converts structured resume JSON to formatted PDF/DOCX
"""

from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Any, cast, Optional
from reportlab.lib import colors
from reportlab.platypus import Flowable
import unicodedata
import re


def sanitize_text(value: str) -> str:
    """Normalize and clean text for safe PDF rendering.

    - Replaces non-breaking hyphen (\u2011) with a regular hyphen.
    - Normalizes unicode characters.
    - Replaces non-breaking spaces with regular spaces.
    - Strips control characters.
    - Fixes whitespace around hyphens in URLs.
    """
    if not isinstance(value, str):
        return ""

    # Specific problematic characters
    value = value.replace("\u2011", "-")  # non-breaking hyphen → regular hyphen

    # Normalize unicode (e.g., smart quotes)
    normalized = unicodedata.normalize("NFKC", value)

    # Replace non-breaking spaces with regular spaces
    normalized = normalized.replace("\u00A0", " ")

    # Strip control characters
    normalized = re.sub(r"[\u0000-\u001F\u007F]", "", normalized)

    # Fix whitespace around hyphens (common in URLs after NFKC normalization)
    # This fixes cases like "linkedin.com/in/name -with-hyphen" → "linkedin.com/in/name-with-hyphen"
    normalized = re.sub(r"\s+-", "-", normalized)
    normalized = re.sub(r"-\s+", "-", normalized)

    return normalized


class HRFlowable(Flowable):
    """Horizontal Line flowable"""
    def __init__(self, width, thickness: float = 1, color=colors.black):
        Flowable.__init__(self)
        self.width = width
        self.thickness = thickness
        self.color = color

    def draw(self):
        """Draw the line"""
        self.canv.setStrokeColor(self.color)
        self.canv.setLineWidth(self.thickness)
        self.canv.line(0, 0, self.width, 0)


class DocumentGenerator:
    """
    Generates professional resume documents from structured JSON data.
    """
    
    @staticmethod
    def extract_text_from_pdf(pdf_file) -> str:
        """
        Extract all text from a PDF file.
        
        Args:
            pdf_file: Uploaded PDF file object
            
        Returns:
            str: Extracted text content
        """
        try:
            pdf_reader = PdfReader(pdf_file)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    @staticmethod
    def create_pdf_from_json(resume_data: Dict, show_summary: bool = True, show_skills: bool = True,
                            show_experience: bool = True, show_projects: bool = True,
                            show_education: bool = True, show_certifications: bool = True,
                            section_order: Optional[List[str]] = None,
                            font_choice: str = "Professional", color_scheme: str = "Blue Professional") -> BytesIO:
        """
        Create a professionally formatted PDF from resume JSON data.
        
        Args:
            resume_data: Structured resume data as dictionary
            show_summary: Whether to include the professional summary section
            show_skills: Whether to include the skills section
            show_experience: Whether to include the experience section
            show_projects: Whether to include the projects section
            show_education: Whether to include the education section
            show_certifications: Whether to include the certifications section
            section_order: List defining the order of sections (default: ['summary', 'education', 'certifications', 'experience', 'skills'])
            
        Returns:
            BytesIO: PDF file as bytes
        """
        # Default section order if not provided
        if section_order is None:
            section_order = ['summary', 'education', 'certifications', 'experience', 'skills']
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )
        
        elements = []
        
        # Define color schemes
        if color_scheme == "Blue Professional":
            primary_color = colors.HexColor('#2c3e50')
            accent_color = colors.HexColor('#3498db')
            line_color = colors.HexColor("#3498db")
        elif color_scheme == "Green Professional":
            primary_color = colors.HexColor('#27ae60')
            accent_color = colors.HexColor('#2ecc71')
            line_color = colors.HexColor("#27ae60")
        else:  # Black & White
            primary_color = colors.black
            accent_color = colors.black
            line_color = colors.HexColor("#cccccc")
        
        # Define font choices
        if font_choice == "Modern":
            name_font = 'Helvetica-Bold'
            header_font = 'Helvetica-Bold'
            body_font = 'Helvetica'
            bullet_font = 'Helvetica'
        elif font_choice == "Classic":
            name_font = 'Times-Bold'
            header_font = 'Times-Bold'
            body_font = 'Times-Roman'
            bullet_font = 'Times-Roman'
        else:  # Professional (default)
            name_font = 'Helvetica-Bold'
            header_font = 'Times-Bold'
            body_font = 'Times-Roman'
            bullet_font = 'Times-Roman'
        
        # Define styles
        name_style = ParagraphStyle(
            'Name',
            fontName=name_font,
            fontSize=22,
            textColor=primary_color,
            alignment=TA_CENTER,
            spaceAfter=6,
            leading=26
        )
        
        contact_style = ParagraphStyle(
            'Contact',
            fontName=body_font,
            fontSize=10,
            alignment=TA_CENTER,
            spaceAfter=10
        )
        
        section_header_style = ParagraphStyle(
            'SectionHeader',
            fontName=header_font,
            fontSize=12,
            textColor=primary_color,
            spaceAfter=6,
            spaceBefore=10,
        )
        
        normal_style = ParagraphStyle(
            'Normal',
            fontName=body_font,
            fontSize=10,
            spaceAfter=4,
            alignment=TA_LEFT,
            leading=14
        )
        
        bullet_style = ParagraphStyle(
            'Bullet',
            fontName=bullet_font,
            fontSize=10,
            spaceAfter=3,
            leftIndent=20,
            bulletIndent=10,
            leading=14
        )
        
        job_title_style = ParagraphStyle(
            'JobTitle',
            fontName=header_font,
            fontSize=11,
            spaceAfter=1
        )
        
        company_style = ParagraphStyle(
            'Company',
            fontName='Times-Italic',
            fontSize=10,
            spaceAfter=4
        )
        
        # Add contact information
        contact = resume_data.get('contact_info', {})
        if contact and contact.get('name'):
            elements.append(Paragraph(sanitize_text(contact['name']), name_style))
        
        # Only include non-empty contact fields
        contact_parts = []
        if contact:
            if contact.get('email') and contact.get('email').strip():
                contact_parts.append(contact['email'].strip())
            if contact.get('phone') and contact.get('phone').strip():
                contact_parts.append(contact['phone'].strip())
            if contact.get('location') and contact.get('location').strip():
                contact_parts.append(contact['location'].strip())
            if contact.get('linkedin') and contact.get('linkedin').strip():
                contact_parts.append(contact['linkedin'].strip())
        
        if contact_parts:
            elements.append(Paragraph(sanitize_text(' | '.join(contact_parts)), contact_style))
        
        # Horizontal Line
        elements.append(Spacer(1, 0.1*inch))
        elements.append(HRFlowable(width=500, thickness=0.5, color=line_color))
        elements.append(Spacer(1, 0.1*inch))

        # Define section rendering functions
        def add_summary_section():
            if show_summary and resume_data.get('professional_summary'):
                elements.append(Paragraph('PROFESSIONAL SUMMARY', section_header_style))
                elements.append(HRFlowable(width=500, thickness=0.2, color=line_color))
                elements.append(Spacer(1, 0.05*inch))
                elements.append(Paragraph(sanitize_text(resume_data['professional_summary']), normal_style))

        def add_skills_section():
            if show_skills and resume_data.get('skills'):
                elements.append(Spacer(1, 0.1*inch))
                elements.append(Paragraph('SKILLS', section_header_style))
                elements.append(HRFlowable(width=500, thickness=0.2, color=line_color))

                skills = resume_data['skills']
                all_skills = []

                if skills.get('technical_skills'):
                    all_skills.extend(skills['technical_skills'])

                if skills.get('tools_technologies'):
                    all_skills.extend(skills['tools_technologies'])

                if skills.get('methodologies'):
                    all_skills.extend(skills['methodologies'])

                if all_skills:
                    elements.append(Paragraph(
                        sanitize_text(', '.join(all_skills)),
                        normal_style
                    ))

        def add_experience_section():
            if show_experience and resume_data.get('experience'):
                elements.append(Spacer(1, 0.1*inch))
                elements.append(Paragraph('PROFESSIONAL EXPERIENCE', section_header_style))
                elements.append(HRFlowable(width=500, thickness=0.2, color=line_color))

                for job in resume_data['experience']:
                    elements.append(Spacer(1, 0.1*inch))
                    # Job title
                    if job.get('job_title'):
                        elements.append(Paragraph(sanitize_text(job['job_title']), job_title_style))

                    # Company, location, and dates
                    company_line_parts = []
                    if job.get('company'):
                        company_line_parts.append(job['company'])
                    if job.get('location'):
                        company_line_parts.append(job['location'])
                    if job.get('dates'):
                        company_line_parts.append(job['dates'])

                    if company_line_parts:
                        elements.append(Paragraph(sanitize_text(' | '.join(company_line_parts)), company_style))

                    # Achievements
                    if job.get('achievements'):
                        for achievement in job['achievements']:
                            elements.append(Paragraph(f"• {sanitize_text(achievement)}", bullet_style))

        def add_projects_section():
            if show_projects and resume_data.get('projects'):
                elements.append(Spacer(1, 0.1*inch))
                elements.append(Paragraph('PROJECTS', section_header_style))
                elements.append(HRFlowable(width=500, thickness=0.2, color=line_color))

                for project in resume_data['projects']:
                    elements.append(Spacer(1, 0.1*inch))
                    # Project name
                    if project.get('project_name'):
                        elements.append(Paragraph(sanitize_text(project['project_name']), job_title_style))

                    # Bullet points (like experience achievements)
                    if project.get('bullets'):
                        for bullet in project['bullets']:
                            elements.append(Paragraph(f"• {sanitize_text(bullet)}", bullet_style))
                    # Fallback to description if bullets don't exist (backward compatibility)
                    elif project.get('description'):
                        elements.append(Paragraph(f"• {sanitize_text(project['description'])}", bullet_style))

        def add_education_section():
            if show_education and resume_data.get('education'):
                elements.append(Spacer(1, 0.1*inch))
                elements.append(Paragraph('EDUCATION', section_header_style))
                elements.append(HRFlowable(width=500, thickness=0.2, color=line_color))

                for edu in resume_data['education']:
                    elements.append(Spacer(1, 0.1*inch))
                    degree_parts = []
                    if edu.get('degree'):
                        degree_parts.append(edu['degree'])
                    if edu.get('institution'):
                        degree_parts.append(edu['institution'])

                    if degree_parts:
                        elements.append(Paragraph(sanitize_text(' - '.join(degree_parts)), job_title_style))

                    edu_details = []
                    if edu.get('location'):
                        edu_details.append(edu['location'])
                    if edu.get('graduation_date'):
                        edu_details.append(edu['graduation_date'])
                    if edu.get('gpa'):
                        edu_details.append(f"GPA: {edu['gpa']}")

                    if edu_details:
                        elements.append(Paragraph(sanitize_text(' | '.join(edu_details)), company_style))

                    if edu.get('relevant_coursework'):
                        elements.append(Paragraph(
                            sanitize_text(f"<b>Relevant Coursework:</b> {', '.join(edu['relevant_coursework'])}"),
                            normal_style
                        ))

        def add_certifications_section():
            if show_certifications and resume_data.get('certifications'):
                elements.append(Spacer(1, 0.1*inch))
                elements.append(Paragraph('CERTIFICATIONS', section_header_style))
                elements.append(HRFlowable(width=500, thickness=0.2, color=line_color))

                for cert in resume_data['certifications']:
                    elements.append(Spacer(1, 0.05*inch))
                    
                    # Format: Certification Name (bold) | Issuing Organization | Date
                    cert_line = ""
                    if cert.get('name'):
                        cert_line = f"<b>{sanitize_text(cert['name'])}</b>"
                    
                    if cert.get('issuer'):
                        cert_line += f" | {sanitize_text(cert['issuer'])}"
                    
                    if cert.get('date'):
                        cert_line += f" | {sanitize_text(cert['date'])}"
                    
                    # Add credential ID if available
                    if cert.get('credential_id'):
                        cert_line += f" | Credential ID: {sanitize_text(cert['credential_id'])}"
                    
                    # Add expiration date if available
                    if cert.get('expiration_date'):
                        cert_line += f" | Expires: {sanitize_text(cert['expiration_date'])}"
                    
                    if cert_line:
                        elements.append(Paragraph('• ' + cert_line, normal_style))
        
        # Map section names to their rendering functions
        section_functions = {
            'summary': add_summary_section,
            'skills': add_skills_section,
            'experience': add_experience_section,
            'projects': add_projects_section,
            'education': add_education_section,
            'certifications': add_certifications_section
        }
        
        # Render sections in the specified order
        for section_name in section_order:
            if section_name in section_functions:
                section_functions[section_name]()
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def create_docx_from_json(resume_data: Dict, show_summary: bool = True, show_skills: bool = True,
                              show_experience: bool = True, show_projects: bool = True,
                              show_education: bool = True, show_certifications: bool = True,
                              section_order: Optional[List[str]] = None) -> BytesIO:
        """
        Create a professionally formatted DOCX from resume JSON data.
        
        Args:
            resume_data: Structured resume data as dictionary
            show_summary: Whether to include the professional summary section
            show_skills: Whether to include the skills section
            show_experience: Whether to include the experience section
            show_projects: Whether to include the projects section
            show_education: Whether to include the education section
            show_certifications: Whether to include the certifications section
            section_order: List defining the order of sections (default: ['summary', 'education', 'certifications', 'experience', 'skills'])
            
        Returns:
            BytesIO: DOCX file as bytes
        """
        # Default section order if not provided
        if section_order is None:
            section_order = ['summary', 'education', 'certifications', 'experience', 'skills']
        buffer = BytesIO()
        doc = Document()
        
        # Set document margins and base font
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Set default font for Normal style
        from docx.styles.style import _ParagraphStyle
        normal_style = cast(_ParagraphStyle, doc.styles['Normal'])
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(10.5)

        # Define colors
        heading_color = RGBColor(0x2c, 0x3e, 0x50)

        # Add contact information
        contact = resume_data.get('contact_info', {})
        if contact and contact.get('name'):
            p = doc.add_paragraph(contact['name'])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = heading_color
        
        # Only include non-empty contact fields
        contact_parts = []
        if contact:
            if contact.get('email') and contact.get('email').strip():
                contact_parts.append(contact['email'].strip())
            if contact.get('phone') and contact.get('phone').strip():
                contact_parts.append(contact['phone'].strip())
            if contact.get('location') and contact.get('location').strip():
                contact_parts.append(contact['location'].strip())
            if contact.get('linkedin') and contact.get('linkedin').strip():
                contact_parts.append(contact['linkedin'].strip())
        
        if contact_parts:
            p = doc.add_paragraph(' | '.join(contact_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            run = p.runs[0]
            run.font.size = Pt(10)

        # --- HR ---
        p = doc.add_paragraph()
        p_border = OxmlElement('w:pBdr')
        p_border.set(qn('w:bottom'), 'w:single w:sz="4" w:space="1" w:color="auto"')
        p._p.get_or_add_pPr().append(p_border)

        def add_section_header(text):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = heading_color
            # Add bottom border
            p_border = OxmlElement('w:pBdr')
            p_border.set(qn('w:bottom'), 'w:single w:sz="4" w:space="1" w:color="auto"')
            p._p.get_or_add_pPr().append(p_border)

        # Define section rendering functions
        def add_summary_section_docx():
            if show_summary and resume_data.get('professional_summary'):
                add_section_header('PROFESSIONAL SUMMARY')
                doc.add_paragraph(resume_data['professional_summary'])

        def add_skills_section_docx():
            if show_skills and resume_data.get('skills'):
                add_section_header('SKILLS')
                skills = resume_data['skills']
                all_skills = []

                if skills.get('technical_skills'):
                    all_skills.extend(skills['technical_skills'])

                if skills.get('tools_technologies'):
                    all_skills.extend(skills['tools_technologies'])

                if skills.get('methodologies'):
                    all_skills.extend(skills['methodologies'])

                if all_skills:
                    p = doc.add_paragraph(', '.join(all_skills))

        def add_experience_section_docx():
            if show_experience and resume_data.get('experience'):
                add_section_header('PROFESSIONAL EXPERIENCE')

                for job in resume_data['experience']:
                    # Job title
                    if job.get('job_title'):
                        p = doc.add_paragraph(job['job_title'])
                        p.paragraph_format.space_before = Pt(8)
                        p.paragraph_format.space_after = Pt(0)
                        run = p.runs[0]
                        run.font.size = Pt(11)
                        run.font.bold = True

                    # Company, location, dates
                    company_parts = []
                    if job.get('company'):
                        company_parts.append(job['company'])
                    if job.get('location'):
                        company_parts.append(job['location'])
                    if job.get('dates'):
                        company_parts.append(job['dates'])

                    if company_parts:
                        p = doc.add_paragraph(' | '.join(company_parts))
                        p.paragraph_format.space_after = Pt(4)
                        run = p.runs[0]
                        run.font.size = Pt(10)
                        run.italic = True

                    # Achievements
                    if job.get('achievements'):
                        for achievement in job['achievements']:
                            p = doc.add_paragraph(achievement, style='List Bullet')
                            p.paragraph_format.left_indent = Inches(0.25)
                            p.paragraph_format.space_after = Pt(2)

        def add_projects_section_docx():
            if show_projects and resume_data.get('projects'):
                add_section_header('PROJECTS')

                for project in resume_data['projects']:
                    # Project name
                    if project.get('project_name'):
                        p = doc.add_paragraph(project['project_name'])
                        p.paragraph_format.space_before = Pt(8)
                        p.paragraph_format.space_after = Pt(2)
                        run = p.runs[0]
                        run.font.size = Pt(11)
                        run.font.bold = True

                    # Bullet points (like experience achievements)
                    if project.get('bullets'):
                        for bullet in project['bullets']:
                            p = doc.add_paragraph(bullet, style='List Bullet')
                            p.paragraph_format.left_indent = Inches(0.25)
                            p.paragraph_format.space_after = Pt(2)
                    # Fallback to description if bullets don't exist (backward compatibility)
                    elif project.get('description'):
                        p = doc.add_paragraph(project['description'], style='List Bullet')
                        p.paragraph_format.left_indent = Inches(0.25)
                        p.paragraph_format.space_after = Pt(6)

        def add_education_section_docx():
            if show_education and resume_data.get('education'):
                add_section_header('EDUCATION')

                for edu in resume_data['education']:
                    degree_parts = []
                    if edu.get('degree'):
                        degree_parts.append(edu['degree'])
                    if edu.get('institution'):
                        degree_parts.append(edu['institution'])

                    if degree_parts:
                        p = doc.add_paragraph(' - '.join(degree_parts))
                        p.paragraph_format.space_before = Pt(8)
                        p.paragraph_format.space_after = Pt(0)
                        run = p.runs[0]
                        run.font.size = Pt(11)
                        run.font.bold = True

                    edu_details = []
                    if edu.get('location'):
                        edu_details.append(edu['location'])
                    if edu.get('graduation_date'):
                        edu_details.append(edu['graduation_date'])
                    if edu.get('gpa'):
                        edu_details.append(f"GPA: {edu['gpa']}")

                    if edu_details:
                        p = doc.add_paragraph(' | '.join(edu_details))
                        run = p.runs[0]
                        run.font.size = Pt(10)

        def add_certifications_section_docx():
            if show_certifications and resume_data.get('certifications'):
                add_section_header('CERTIFICATIONS')

                for cert in resume_data['certifications']:
                    # Create a paragraph for the certification
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(4)
                    
                    # Add certification name in bold
                    if cert.get('name'):
                        run = p.add_run(cert['name'])
                        run.bold = True
                    
                    # Add issuer
                    if cert.get('issuer'):
                        p.add_run(f" | {cert['issuer']}")
                    
                    # Add date
                    if cert.get('date'):
                        p.add_run(f" | {cert['date']}")
                    
                    # Add credential ID if available
                    if cert.get('credential_id'):
                        run = p.add_run(f"\n    Credential ID: {cert['credential_id']}")
                        run.font.size = Pt(9)
                        run.italic = True
                    
                    # Add expiration date if available
                    if cert.get('expiration_date'):
                        run = p.add_run(f"\n    Expires: {cert['expiration_date']}")
                        run.font.size = Pt(9)
                        run.italic = True
        
        # Map section names to their rendering functions
        section_functions_docx = {
            'summary': add_summary_section_docx,
            'skills': add_skills_section_docx,
            'experience': add_experience_section_docx,
            'projects': add_projects_section_docx,
            'education': add_education_section_docx,
            'certifications': add_certifications_section_docx
        }
        
        # Render sections in the specified order
        for section_name in section_order:
            if section_name in section_functions_docx:
                section_functions_docx[section_name]()
        
        # Save to buffer
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def create_cover_letter_pdf(cover_letter_text: str, candidate_name: Optional[str] = None,
                                 company_name: Optional[str] = None, role_title: Optional[str] = None) -> BytesIO:
        """
        Create a professionally formatted PDF for a cover letter.
        
        Args:
            cover_letter_text: The cover letter content
            candidate_name: Optional candidate name for header
            company_name: Optional company name for context
            role_title: Optional role title for context
            
        Returns:
            BytesIO: PDF file as bytes
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        elements = []
        
        # Define styles for cover letter
        name_style = ParagraphStyle(
            'Name',
            fontName='Times-Bold',
            fontSize=14,
            textColor=colors.HexColor('#2c3e50'),
            alignment=TA_LEFT,
            spaceAfter=12,
            leading=18
        )
        
        body_style = ParagraphStyle(
            'Body',
            fontName='Times-Roman',
            fontSize=11,
            alignment=TA_LEFT,
            spaceAfter=12,
            leading=16,
            firstLineIndent=0
        )
        
        # Add candidate name if provided
        if candidate_name:
            elements.append(Paragraph(sanitize_text(candidate_name), name_style))
        
        # Add some space before the letter body
        elements.append(Spacer(1, 0.2*inch))
        
        # Split cover letter into paragraphs and add them
        paragraphs = cover_letter_text.strip().split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Clean up the paragraph text
                cleaned_para = sanitize_text(para.strip())
                elements.append(Paragraph(cleaned_para, body_style))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def create_qa_text_file(recruiter_qa: List[Dict[str, str]]) -> str:
        """
        Create a formatted text file content from Q&A data.
        
        Args:
            recruiter_qa: List of dictionaries with 'question' and 'answer' keys
            
        Returns:
            str: Formatted text content
        """
        if not recruiter_qa:
            return "No questions and answers available."
        
        lines = ["RECRUITER QUESTIONS & ANSWERS", "=" * 50, ""]
        
        for idx, qa in enumerate(recruiter_qa, start=1):
            question = qa.get('question', '').strip()
            answer = qa.get('answer', '').strip()
            
            lines.append(f"Question {idx}:")
            lines.append(question)
            lines.append("")
            lines.append(f"Answer {idx}:")
            lines.append(answer)
            lines.append("")
            lines.append("-" * 50)
            lines.append("")
        
        return "\n".join(lines)
